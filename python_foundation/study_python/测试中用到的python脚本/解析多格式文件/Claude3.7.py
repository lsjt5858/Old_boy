#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Author: 熊🐻来个🥬
# @Date:  2025/3/3
# @Description: [对文件功能等的简要描述（可自行添加）]
import sys
import json
import yaml
import pandas as pd
import markdown
import base64
import requests
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QPushButton, QFileDialog, QTabWidget, QTextEdit,
                             QComboBox, QLineEdit, QMessageBox, QGroupBox, QScrollArea,
                             QSplitter, QTreeWidget, QTreeWidgetItem)
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtGui import QFont, QIcon
import docx
from docx.shared import Pt
import re
import io
import html2text


class MultiFormatTestTool(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()
        self.current_file_path = None
        self.current_file_content = None
        self.current_file_type = None
        self.parsed_data = None

    def initUI(self):
        self.setWindowTitle('多格式文档支持的智能测试用例生成工具')
        self.setGeometry(100, 100, 1200, 800)

        # 创建主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)

        # 主布局
        main_layout = QVBoxLayout()
        main_widget.setLayout(main_layout)

        # 顶部操作区
        top_group = QGroupBox("文档导入")
        top_layout = QHBoxLayout()
        top_group.setLayout(top_layout)

        self.import_btn = QPushButton("导入文档")
        self.import_btn.clicked.connect(self.import_document)
        self.file_path_display = QLineEdit()
        self.file_path_display.setReadOnly(True)
        self.file_type_combo = QComboBox()
        self.file_type_combo.addItems(
            ["HAR", "Swagger JSON", "Swagger YAML", "Excel", "Markdown", "JSON", "YAML", "TXT"])

        top_layout.addWidget(self.import_btn)
        top_layout.addWidget(self.file_path_display)
        top_layout.addWidget(QLabel("文件类型:"))
        top_layout.addWidget(self.file_type_combo)

        main_layout.addWidget(top_group)

        # 分割器
        splitter = QSplitter(Qt.Vertical)

        # 中间内容区域
        content_tabs = QTabWidget()

        # 原始内容预览
        self.raw_content_preview = QTextEdit()
        self.raw_content_preview.setReadOnly(True)
        content_tabs.addTab(self.raw_content_preview, "原始内容预览")

        # 解析结果预览
        self.parsed_tree = QTreeWidget()
        self.parsed_tree.setHeaderLabels(["接口信息", "详情"])
        self.parsed_tree.setColumnWidth(0, 300)
        content_tabs.addTab(self.parsed_tree, "解析结果")

        splitter.addWidget(content_tabs)

        # 底部操作区域
        bottom_group = QGroupBox("测试用例生成")
        bottom_layout = QVBoxLayout()
        bottom_group.setLayout(bottom_layout)

        prompt_layout = QHBoxLayout()
        prompt_layout.addWidget(QLabel("提示词:"))
        self.prompt_input = QLineEdit()
        self.prompt_input.setPlaceholderText("请输入提示词以指导测试用例生成...")
        prompt_layout.addWidget(self.prompt_input)
        self.generate_btn = QPushButton("生成测试用例")
        self.generate_btn.clicked.connect(self.generate_test_cases)
        prompt_layout.addWidget(self.generate_btn)

        bottom_layout.addLayout(prompt_layout)

        # 测试用例预览
        self.test_case_preview = QTextEdit()
        bottom_layout.addWidget(self.test_case_preview)

        # 导出选项
        export_layout = QHBoxLayout()
        export_layout.addWidget(QLabel("导出格式:"))
        self.export_format_combo = QComboBox()
        self.export_format_combo.addItems(["JSON", "YAML", "Excel", "Markdown", "Docx"])
        export_layout.addWidget(self.export_format_combo)
        self.export_btn = QPushButton("导出测试用例")
        self.export_btn.clicked.connect(self.export_test_cases)
        export_layout.addWidget(self.export_btn)

        bottom_layout.addLayout(export_layout)

        splitter.addWidget(bottom_group)

        # 设置分割器比例
        splitter.setSizes([400, 400])

        main_layout.addWidget(splitter)

        # 设置样式
        self.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 1px solid #ccc;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QPushButton {
                background-color: #4a86e8;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #3a76d8;
            }
            QTextEdit, QTreeWidget {
                border: 1px solid #ddd;
                border-radius: 3px;
            }
        """)

    def import_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择文件",
            "",
            "所有文件 (*.*);;HAR文件 (*.har);;JSON文件 (*.json);;YAML文件 (*.yaml *.yml);;Excel文件 (*.xlsx *.xls);;Markdown文件 (*.md);;文本文件 (*.txt)"
        )

        if file_path:
            self.current_file_path = file_path
            self.file_path_display.setText(file_path)

            # 根据文件扩展名自动选择文件类型
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.har':
                self.file_type_combo.setCurrentText("HAR")
            elif ext == '.json':
                # 需要判断是普通JSON还是Swagger JSON
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = json.load(f)
                    if 'swagger' in content or 'openapi' in content:
                        self.file_type_combo.setCurrentText("Swagger JSON")
                    else:
                        self.file_type_combo.setCurrentText("JSON")
                except:
                    self.file_type_combo.setCurrentText("JSON")
            elif ext in ['.yaml', '.yml']:
                # 需要判断是普通YAML还是Swagger YAML
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = yaml.safe_load(f)
                    if content and ('swagger' in content or 'openapi' in content):
                        self.file_type_combo.setCurrentText("Swagger YAML")
                    else:
                        self.file_type_combo.setCurrentText("YAML")
                except:
                    self.file_type_combo.setCurrentText("YAML")
            elif ext in ['.xlsx', '.xls']:
                self.file_type_combo.setCurrentText("Excel")
            elif ext == '.md':
                self.file_type_combo.setCurrentText("Markdown")
            elif ext == '.txt':
                self.file_type_combo.setCurrentText("TXT")

            self.load_document(file_path)

    def load_document(self, file_path):
        try:
            file_type = self.file_type_combo.currentText()
            self.current_file_type = file_type

            content = None
            if file_type == "HAR":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = json.load(f)
                    self.raw_content_preview.setText(json.dumps(content, indent=2, ensure_ascii=False))

            elif file_type in ["Swagger JSON", "JSON"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = json.load(f)
                    self.raw_content_preview.setText(json.dumps(content, indent=2, ensure_ascii=False))

            elif file_type in ["Swagger YAML", "YAML"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = yaml.safe_load(f)
                    self.raw_content_preview.setText(yaml.dump(content, allow_unicode=True))

            elif file_type == "Excel":
                df = pd.read_excel(file_path)
                content = df
                self.raw_content_preview.setText(str(df))

            elif file_type == "Markdown":
                with open(file_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                    content = md_content
                    # 同时展示原始markdown和解析后的HTML
                    self.raw_content_preview.setText(md_content)

            elif file_type == "TXT":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    self.raw_content_preview.setText(content)

            self.current_file_content = content

            # 解析文档内容
            self.parse_document()

        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载文档失败: {str(e)}")

    def parse_document(self):
        try:
            # 清空之前的解析结果
            self.parsed_tree.clear()

            file_type = self.current_file_type
            content = self.current_file_content

            if file_type == "HAR":
                self.parse_har(content)
            elif file_type == "Swagger JSON" or file_type == "Swagger YAML":
                self.parse_swagger(content)
            elif file_type == "Excel":
                self.parse_excel(content)
            elif file_type == "Markdown":
                self.parse_markdown(content)
            elif file_type == "JSON":
                self.parse_json(content)
            elif file_type == "YAML":
                self.parse_yaml(content)
            elif file_type == "TXT":
                self.parse_txt(content)

        except Exception as e:
            QMessageBox.warning(self, "警告", f"解析文档失败: {str(e)}")

    def parse_har(self, content):
        try:
            # 提取HAR文件中的请求信息
            entries = content.get('log', {}).get('entries', [])

            self.parsed_data = []

            for entry in entries:
                request = entry.get('request', {})
                response = entry.get('response', {})

                # 提取基本信息
                method = request.get('method', '')
                url = request.get('url', '')
                status = response.get('status', '')
                status_text = response.get('statusText', '')

                # 提取请求头
                headers = {}
                for header in request.get('headers', []):
                    headers[header.get('name')] = header.get('value')

                # 提取请求参数
                params = {}
                for param in request.get('queryString', []):
                    params[param.get('name')] = param.get('value')

                # 提取请求体
                post_data = request.get('postData', {})
                body = post_data.get('text', '')

                # 提取响应内容
                response_content = response.get('content', {})
                response_body = response_content.get('text', '')

                # 保存解析结果
                api_info = {
                    'method': method,
                    'url': url,
                    'headers': headers,
                    'params': params,
                    'body': body,
                    'status': status,
                    'status_text': status_text,
                    'response': response_body
                }

                self.parsed_data.append(api_info)

                # 添加到树形控件
                root_item = QTreeWidgetItem(self.parsed_tree)
                root_item.setText(0, f"{method} {url}")
                root_item.setText(1, f"Status: {status} {status_text}")

                headers_item = QTreeWidgetItem(root_item)
                headers_item.setText(0, "Headers")
                for key, value in headers.items():
                    header_item = QTreeWidgetItem(headers_item)
                    header_item.setText(0, key)
                    header_item.setText(1, value)

                params_item = QTreeWidgetItem(root_item)
                params_item.setText(0, "Parameters")
                for key, value in params.items():
                    param_item = QTreeWidgetItem(params_item)
                    param_item.setText(0, key)
                    param_item.setText(1, value)

                body_item = QTreeWidgetItem(root_item)
                body_item.setText(0, "Request Body")
                body_item.setText(1, body[:100] + "..." if len(body) > 100 else body)

                response_item = QTreeWidgetItem(root_item)
                response_item.setText(0, "Response")
                response_item.setText(1, response_body[:100] + "..." if len(response_body) > 100 else response_body)

            self.parsed_tree.expandToDepth(0)

        except Exception as e:
            raise Exception(f"解析HAR文件失败: {str(e)}")

    def parse_swagger(self, content):
        try:
            # 处理Swagger/OpenAPI文档
            swagger_version = content.get('swagger') or content.get('openapi')
            info = content.get('info', {})
            title = info.get('title', 'API文档')

            # 获取基础路径
            base_path = content.get('basePath', '')

            # 获取所有路径
            paths = content.get('paths', {})

            self.parsed_data = []

            # 添加基本信息
            root_item = QTreeWidgetItem(self.parsed_tree)
            root_item.setText(0, title)
            root_item.setText(1, f"Swagger/OpenAPI版本: {swagger_version}")

            # 添加路径信息
            paths_item = QTreeWidgetItem(root_item)
            paths_item.setText(0, "API路径")

            for path, path_info in paths.items():
                path_item = QTreeWidgetItem(paths_item)
                path_item.setText(0, path)

                for method, method_info in path_info.items():
                    if method not in ['get', 'post', 'put', 'delete', 'patch', 'options', 'head']:
                        continue

                    method_item = QTreeWidgetItem(path_item)
                    method_item.setText(0, method.upper())
                    method_item.setText(1, method_info.get('summary', ''))

                    # 处理参数
                    parameters = method_info.get('parameters', [])
                    if parameters:
                        params_item = QTreeWidgetItem(method_item)
                        params_item.setText(0, "Parameters")

                        for param in parameters:
                            param_item = QTreeWidgetItem(params_item)
                            param_item.setText(0, param.get('name', ''))
                            param_item.setText(1, f"({param.get('in', '')}): {param.get('description', '')}")

                    # 处理响应
                    responses = method_info.get('responses', {})
                    if responses:
                        resp_item = QTreeWidgetItem(method_item)
                        resp_item.setText(0, "Responses")

                        for status, resp_info in responses.items():
                            status_item = QTreeWidgetItem(resp_item)
                            status_item.setText(0, status)
                            status_item.setText(1, resp_info.get('description', ''))

                    # 保存解析结果
                    api_info = {
                        'method': method.upper(),
                        'path': path,
                        'full_path': f"{base_path}{path}",
                        'summary': method_info.get('summary', ''),
                        'description': method_info.get('description', ''),
                        'parameters': parameters,
                        'responses': responses
                    }

                    self.parsed_data.append(api_info)

            self.parsed_tree.expandToDepth(1)

        except Exception as e:
            raise Exception(f"解析Swagger文档失败: {str(e)}")

    def parse_excel(self, content):
        try:
            # 假设Excel中包含API信息，需要根据实际表格结构调整
            df = content

            self.parsed_data = []

            # 尝试检测Excel文件结构
            if all(col in df.columns for col in ['Method', 'URL', 'Headers', 'Body']):
                # 标准API定义格式
                for index, row in df.iterrows():
                    method = row.get('Method', '')
                    url = row.get('URL', '')
                    headers = row.get('Headers', '')
                    body = row.get('Body', '')

                    if isinstance(headers, str):
                        try:
                            headers = json.loads(headers)
                        except:
                            headers = {}

                    if isinstance(body, str):
                        try:
                            body = json.loads(body)
                        except:
                            pass

                    # 保存解析结果
                    api_info = {
                        'method': method,
                        'url': url,
                        'headers': headers,
                        'body': body
                    }

                    self.parsed_data.append(api_info)

                    # 添加到树形控件
                    root_item = QTreeWidgetItem(self.parsed_tree)
                    root_item.setText(0, f"{method} {url}")

                    headers_item = QTreeWidgetItem(root_item)
                    headers_item.setText(0, "Headers")
                    if isinstance(headers, dict):
                        for key, value in headers.items():
                            header_item = QTreeWidgetItem(headers_item)
                            header_item.setText(0, key)
                            header_item.setText(1, str(value))

                    body_item = QTreeWidgetItem(root_item)
                    body_item.setText(0, "Body")
                    if isinstance(body, dict):
                        for key, value in body.items():
                            body_param_item = QTreeWidgetItem(body_item)
                            body_param_item.setText(0, key)
                            body_param_item.setText(1, str(value))
                    else:
                        body_item.setText(1, str(body))
            else:
                # 通用表格解析
                # 将每一列作为一个节点
                root_item = QTreeWidgetItem(self.parsed_tree)
                root_item.setText(0, "Excel数据")

                for col in df.columns:
                    col_item = QTreeWidgetItem(root_item)
                    col_item.setText(0, col)

                    # 显示前10行数据
                    for i in range(min(10, len(df))):
                        row_item = QTreeWidgetItem(col_item)
                        row_item.setText(0, f"Row {i + 1}")
                        row_item.setText(1, str(df.iloc[i][col]))

                # 尝试解析数据结构
                self.parsed_data = df.to_dict('records')

            self.parsed_tree.expandToDepth(0)

        except Exception as e:
            raise Exception(f"解析Excel文件失败: {str(e)}")

    def parse_markdown(self, content):
        try:
            md_content = content

            # 使用正则表达式提取API信息
            api_pattern = r'#{1,3}\s+(.*?)\s*\n(.*?)(?=#{1,3}|\Z)'
            api_matches = re.findall(api_pattern, md_content, re.DOTALL)

            # 匹配URL和方法
            url_method_pattern = r'([A-Z]+)\s+(\/\S*)'

            # 匹配请求头、参数和响应
            headers_pattern = r'[Hh]eaders?:?\s*```(?:json)?\s*(.*?)\s*```'
            params_pattern = r'[Pp]arameters?:?\s*```(?:json)?\s*(.*?)\s*```'
            body_pattern = r'[Bb]ody:?\s*```(?:json)?\s*(.*?)\s*```'
            response_pattern = r'[Rr]esponse:?\s*```(?:json)?\s*(.*?)\s*```'

            self.parsed_data = []

            for title, content in api_matches:
                # 尝试提取URL和方法
                url_method_match = re.search(url_method_pattern, title + "\n" + content)
                method = ''
                url = ''

                if url_method_match:
                    method = url_method_match.group(1)
                    url = url_method_match.group(2)

                # 提取请求头
                headers_match = re.search(headers_pattern, content, re.DOTALL)
                headers = {}
                if headers_match:
                    try:
                        headers_text = headers_match.group(1)
                        headers = json.loads(headers_text)
                    except:
                        pass

                # 提取参数
                params_match = re.search(params_pattern, content, re.DOTALL)
                params = {}
                if params_match:
                    try:
                        params_text = params_match.group(1)
                        params = json.loads(params_text)
                    except:
                        pass

                # 提取请求体
                body_match = re.search(body_pattern, content, re.DOTALL)
                body = ''
                if body_match:
                    body = body_match.group(1)

                # 提取响应
                response_match = re.search(response_pattern, content, re.DOTALL)
                response = ''
                if response_match:
                    response = response_match.group(1)

                # 保存解析结果
                api_info = {
                    'title': title,
                    'method': method,
                    'url': url,
                    'headers': headers,
                    'params': params,
                    'body': body,
                    'response': response
                }

                self.parsed_data.append(api_info)

                # 添加到树形控件
                root_item = QTreeWidgetItem(self.parsed_tree)
                root_item.setText(0, title)
                if method and url:
                    root_item.setText(1, f"{method} {url}")

                if headers:
                    headers_item = QTreeWidgetItem(root_item)
                    headers_item.setText(0, "Headers")
                    for key, value in headers.items():
                        header_item = QTreeWidgetItem(headers_item)
                        header_item.setText(0, key)
                        header_item.setText(1, str(value))

                if params:
                    params_item = QTreeWidgetItem(root_item)
                    params_item.setText(0, "Parameters")
                    for key, value in params.items():
                        param_item = QTreeWidgetItem(params_item)
                        param_item.setText(0, key)
                        param_item.setText(1, str(value))

                if body:
                    body_item = QTreeWidgetItem(root_item)
                    body_item.setText(0, "Request Body")
                    body_item.setText(1, body[:100] + "..." if len(body) > 100 else body)

                if response:
                    response_item = QTreeWidgetItem(root_item)
                    response_item.setText(0, "Response")
                    response_item.setText(1, response[:100] + "..." if len(response) > 100 else response)

            self.parsed_tree.expandToDepth(0)

        except Exception as e:
            raise Exception(f"解析Markdown文件失败: {str(e)}")

    def parse_json(self, content):
        try:
            # 检查是否为API接口集合
            if isinstance(content, list) and all(isinstance(item, dict) for item in content):
                # 可能是API集合
                self.parsed_data = []

                for item in content:
                    # 尝试识别常见的API字段
                    method = item.get('method', '')
                    url = item.get('url', '') or item.get('path', '')
                    headers = item.get('headers', {})
                    params = item.get('params', {}) or item.get('parameters', {})
                    body = item.get('body', '') or item.get('data', '')

                    api_info = {
                        'method': method,
                        'url': url,
                        'headers': headers,
                        'params': params,
                        'body': body
                    }

                    self.parsed_data.append(api_info)

                    # 添加到树形控件
                    root_item = QTreeWidgetItem(self.parsed_tree)
                    if method and url:
                        root_item.setText(0, f"{method} {url}")
                    else:
                        root_item.setText(0, f"Item {len(self.parsed_data)}")

                    # 添加详细信息
                    for key, value in item.items():
                        if isinstance(value, dict):
                            sub_item = QTreeWidgetItem(root_item)
                            sub_item.setText(0, key)

                            for sub_key, sub_value in value.items():
                                sub_sub_item = QTreeWidgetItem(sub_item)
                                sub_sub_item.setText(0, sub_key)
                                sub_sub_item.setText(1, str(sub_value)[:100])
                        else:
                            sub_item = QTreeWidgetItem(root_item)
                            sub_item.setText(0, key)
                            sub_item.setText(1, str(value)[:100])
            else:
                # 通用JSON解析
                self.parsed_data = content

                def add_json_to_tree(json_obj, parent_item):
                    if isinstance(json_obj, dict):
                        for key, value in json_obj.items():
                            item = QTreeWidgetItem(parent_item)
                            item.setText(0, str(key))

                            if isinstance(value, (dict, list)):
                                add_json_to_tree(value, item)
                            else:
                                item.setText(1, str(value)[:100])
                    elif isinstance(json_obj, list):
                        for i, value in enumerate(json_obj):
                            item = QTreeWidgetItem(parent_item)
                            item.setText(0, f"[{i}]")

                            if isinstance(value, (dict, list)):
                                add_json_to_tree(value, item)
                            else:
                                item.setText(1, str(value)[:100])

                root_item = QTreeWidgetItem(self.parsed_tree)
                root_item.setText(0, "JSON数据")

                add_json_to_tree(content, root_item)

            self.parsed_tree.expandToDepth(1)

        except Exception as e:
            raise Exception(f"解析JSON文件失败: {str(e)}")

    def parse_yaml(self, content):
        try:
            # 将YAML解析为JSON再处理
            json_content = content  # YAML已经在加载时转换为Python对象
            self.parse_json(json_content)

        except Exception as e:
            raise Exception(f"解析YAML文件失败: {str(e)}")

    def parse_txt(self, content):
        try:
            # 尝试从文本文件中提取API信息

            # 查找常见的API模式
            url_pattern = r'(GET|POST|PUT|DELETE|PATCH)\s+(https?://\S+)'
            header_pattern = r'Header(?:s)?[\s\:]+(.+?)(?=\n\n|\Z)'
            body_pattern = r'Body[\s\:]+(.+?)(?=\n\n|\Z)'

            url_matches = re.findall(url_pattern, content)

            self.parsed_data = []

            if url_matches:
                # 可能包含API信息
                for method, url in url_matches:
                    # 尝试提取对应的请求部分
                    request_pattern = f"{method}\\s+{re.escape(url)}(.+?)(?=(?:GET|POST|PUT|DELETE|PATCH)\\s+https?://|\\Z)"
                    request_match = re.search(request_pattern, content, re.DOTALL)

                    headers = {}
                    body = ""

                    if request_match:
                        request_content = request_match.group(1)

                        # 提取请求头
                        header_match = re.search(header_pattern, request_content, re.DOTALL)
                        if header_match:
                            header_content = header_match.group(1).strip()
                            header_lines = header_content.split('\n')
                            for line in header_lines:
                                if ':' in line:
                                    key, value = line.split(':', 1)
                                    headers[key.strip()] = value.strip()

                        # 提取请求体
                        body_match = re.search(body_pattern, request_content, re.DOTALL)
                        if body_match:
                            body = body_match.group(1).strip()

                    # 保存解析结果
                    api_info = {
                        'method': method,
                        'url': url,
                        'headers': headers,
                        'body': body
                    }

                    self.parsed_data.append(api_info)

                    # 添加到树形控件
                    root_item = QTreeWidgetItem(self.parsed_tree)
                    root_item.setText(0, f"{method} {url}")

                    # 添加请求头
                    if headers:
                        headers_item = QTreeWidgetItem(root_item)
                        headers_item.setText(0, "Headers")
                        for key, value in headers.items():
                            header_item = QTreeWidgetItem(headers_item)
                            header_item.setText(0, key)
                            header_item.setText(1, value)

                    # 添加请求体
                    if body:
                        body_item = QTreeWidgetItem(root_item)
                        body_item.setText(0, "Request Body")
                        body_item.setText(1, body[:100] + "..." if len(body) > 100 else body)
            else:
                # 通用文本解析
                lines = content.split('\n')
                root_item = QTreeWidgetItem(self.parsed_tree)
                root_item.setText(0, "Text Content")

                for i, line in enumerate(lines[:100]):  # 只显示前100行
                    line_item = QTreeWidgetItem(root_item)
                    line_item.setText(0, f"Line {i + 1}")
                    line_item.setText(1, line)

                # 保存文本内容
                self.parsed_data = content

            self.parsed_tree.expandToDepth(0)

        except Exception as e:
            raise Exception(f"解析TXT文件失败: {str(e)}")

    def generate_test_cases(self):
        try:
            if not self.parsed_data:
                QMessageBox.warning(self, "警告", "请先导入并解析文档")
                return

            prompt = self.prompt_input.text()
            if not prompt:
                prompt = "为以下API生成完整的接口测试用例，包括正向测试和异常测试"

            # 准备调用百炼API的数据
            api_data = self.prepare_api_data()

            # 调用阿里云百炼 DeepSeek API
            test_cases = self.call_deepseek_api(prompt, api_data)

            # 显示生成的测试用例
            self.test_case_preview.setText(test_cases)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成测试用例失败: {str(e)}")

    def prepare_api_data(self):
        # 根据不同的文件类型，准备API数据
        file_type = self.current_file_type
        data = self.parsed_data

        # 格式化API数据为易于模型理解的格式
        formatted_data = {}

        if file_type == "HAR":
            formatted_data = {
                "apis": []
            }

            for api in data:
                formatted_api = {
                    "method": api.get('method', ''),
                    "url": api.get('url', ''),
                    "headers": api.get('headers', {}),
                    "query_params": api.get('params', {}),
                    "body": api.get('body', ''),
                    "response": {
                        "status": api.get('status', ''),
                        "body": api.get('response', '')
                    }
                }
                formatted_data["apis"].append(formatted_api)

        elif file_type in ["Swagger JSON", "Swagger YAML"]:
            formatted_data = {
                "apis": []
            }

            for api in data:
                formatted_api = {
                    "method": api.get('method', ''),
                    "path": api.get('path', ''),
                    "full_path": api.get('full_path', ''),
                    "summary": api.get('summary', ''),
                    "description": api.get('description', ''),
                    "parameters": api.get('parameters', []),
                    "responses": api.get('responses', {})
                }
                formatted_data["apis"].append(formatted_api)

        elif file_type == "Excel":
            if isinstance(data, list) and all(isinstance(item, dict) for item in data):
                formatted_data = {"apis": data}
            else:
                formatted_data = {"data": data}

        elif file_type == "Markdown":
            formatted_data = {
                "apis": []
            }

            for api in data:
                formatted_api = {
                    "title": api.get('title', ''),
                    "method": api.get('method', ''),
                    "url": api.get('url', ''),
                    "headers": api.get('headers', {}),
                    "params": api.get('params', {}),
                    "body": api.get('body', ''),
                    "response": api.get('response', '')
                }
                formatted_data["apis"].append(formatted_api)

        elif file_type in ["JSON", "YAML"]:
            if isinstance(data, list) and all(isinstance(item, dict) for item in data):
                formatted_data = {"apis": data}
            else:
                formatted_data = {"data": data}

        elif file_type == "TXT":
            if isinstance(data, list) and all(isinstance(item, dict) for item in data):
                formatted_data = {"apis": data}
            else:
                formatted_data = {"text": data}

        return json.dumps(formatted_data, ensure_ascii=False)

    def call_deepseek_api(self, prompt, api_data):
        try:
            # 这里是一个模拟，实际实现中需要替换为真实的百炼API调用
            # 以下是模拟的调用过程，实际项目中需要替换为真实的API调用

            # 模拟响应
            import time
            time.sleep(1)  # 模拟API调用延迟

            # 根据API数据生成测试用例
            if "apis" in json.loads(api_data):
                apis = json.loads(api_data)["apis"]
                test_cases = self.generate_mock_test_cases(prompt, apis)
            else:
                test_cases = "无法从当前数据生成测试用例，请检查文档格式或提供更详细的API信息。"

            return test_cases

            # 实际实现可能类似下面的代码
            """
            endpoint = "https://bailian-api.aliyun.com/v1/llm/generation"
            api_key = "您的API密钥"

            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }

            data = {
                "model": "deepseek-coder",
                "messages": [
                    {"role": "system", "content": "你是一个专业的接口测试用例生成专家，根据提供的API信息生成完整的接口测试用例。"},
                    {"role": "user", "content": f"{prompt}\n\n以下是API信息:\n{api_data}"}
                ],
                "temperature": 0.7,
                "max_tokens": 4000
            }
            response = requests.post(endpoint, headers=headers, json=data)

            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"]
            else:
                raise Exception(f"API调用失败: {response.status_code}, {response.text}")
            """

        except Exception as e:
            raise Exception(f"调用DeepSeek API失败: {str(e)}")

    def generate_mock_test_cases(self, prompt, apis):
        # 这是一个模拟函数，生成示例测试用例
        test_cases = f"# 接口测试用例\n\n根据提示词: \"{prompt}\"\n\n"

        for i, api in enumerate(apis[:3]):  # 只处理前3个API作为示例
            method = api.get("method", "")
            url = api.get("url", "") or api.get("path", "")

            test_cases += f"## API {i + 1}: {method} {url}\n\n"

            # 正向测试用例
            test_cases += "### 正向测试用例\n\n"
            test_cases += "```python\n"
            test_cases += f"import requests\nimport pytest\n\n"
            test_cases += f"def test_{method.lower()}_success():\n"
            test_cases += f"    # 测试目标: 验证API在正确参数下的成功响应\n"

            # 构建请求参数
            headers = api.get("headers", {})
            params = api.get("params", {}) or api.get("query_params", {})
            body = api.get("body", "")

            # 构建请求代码
            test_cases += f"    url = '{url}'\n"

            if headers:
                test_cases += f"    headers = {json.dumps(headers, indent=4, ensure_ascii=False)}\n"
            else:
                test_cases += f"    headers = {{'Content-Type': 'application/json'}}\n"

            if params:
                test_cases += f"    params = {json.dumps(params, indent=4, ensure_ascii=False)}\n"

            if body:
                if isinstance(body, str) and body.strip():
                    try:
                        body_json = json.loads(body)
                        test_cases += f"    data = {json.dumps(body_json, indent=4, ensure_ascii=False)}\n"
                    except:
                        test_cases += f"    data = {repr(body)}\n"
                elif isinstance(body, dict):
                    test_cases += f"    data = {json.dumps(body, indent=4, ensure_ascii=False)}\n"

            # 发送请求
            if method.upper() in ["GET"]:
                test_cases += f"    response = requests.get(url, headers=headers, params=params)\n\n"
            elif method.upper() in ["POST"]:
                test_cases += f"    response = requests.post(url, headers=headers, json=data)\n\n"
            elif method.upper() in ["PUT"]:
                test_cases += f"    response = requests.put(url, headers=headers, json=data)\n\n"
            elif method.upper() in ["DELETE"]:
                test_cases += f"    response = requests.delete(url, headers=headers)\n\n"
            else:
                test_cases += f"    response = requests.request('{method}', url, headers=headers)\n\n"

            # 断言
            test_cases += f"    # 断言响应状态码为200\n"
            test_cases += f"    assert response.status_code == 200\n"
            test_cases += f"    # 断言响应包含预期字段\n"
            test_cases += f"    assert 'data' in response.json()\n"
            test_cases += f"    # 其他业务相关断言\n"

            test_cases += "```\n\n"

            # 异常测试用例
            test_cases += "### 异常测试用例\n\n"
            test_cases += "```python\n"
            test_cases += f"def test_{method.lower()}_invalid_parameter():\n"
            test_cases += f"    # 测试目标: 验证API在参数错误时的错误处理\n"

            # 构建请求代码（使用无效参数）
            test_cases += f"    url = '{url}'\n"
            test_cases += f"    headers = {{'Content-Type': 'application/json'}}\n"

            if method.upper() in ["GET"]:
                test_cases += f"    params = {{'invalid_param': 'invalid_value'}}\n"
                test_cases += f"    response = requests.get(url, headers=headers, params=params)\n\n"
            elif method.upper() in ["POST"]:
                test_cases += f"    data = {{'invalid_field': 'invalid_value'}}\n"
                test_cases += f"    response = requests.post(url, headers=headers, json=data)\n\n"
            elif method.upper() in ["PUT"]:
                test_cases += f"    data = {{'invalid_field': 'invalid_value'}}\n"
                test_cases += f"    response = requests.put(url, headers=headers, json=data)\n\n"
            elif method.upper() in ["DELETE"]:
                test_cases += f"    params = {{'invalid_param': 'invalid_value'}}\n"
                test_cases += f"    response = requests.delete(url, headers=headers, params=params)\n\n"
            else:
                test_cases += f"    data = {{'invalid_field': 'invalid_value'}}\n"
                test_cases += f"    response = requests.request('{method}', url, headers=headers, json=data)\n\n"

            # 断言
            test_cases += f"    # 断言响应状态码为400（参数错误）\n"
            test_cases += f"    assert response.status_code == 400\n"
            test_cases += f"    # 断言错误信息\n"
            test_cases += f"    assert 'error' in response.json()\n"

            test_cases += "```\n\n"

            # 认证测试用例
            test_cases += "### 认证测试用例\n\n"
            test_cases += "```python\n"
            test_cases += f"def test_{method.lower()}_unauthorized():\n"
            test_cases += f"    # 测试目标: 验证API在未授权时的错误处理\n"

            # 构建请求代码（无授权信息）
            test_cases += f"    url = '{url}'\n"
            test_cases += f"    # 不提供认证信息\n"
            test_cases += f"    headers = {{'Content-Type': 'application/json'}}\n"

            if method.upper() in ["GET"]:
                test_cases += f"    response = requests.get(url, headers=headers)\n\n"
            elif method.upper() in ["POST"]:
                test_cases += f"    response = requests.post(url, headers=headers)\n\n"
            elif method.upper() in ["PUT"]:
                test_cases += f"    response = requests.put(url, headers=headers)\n\n"
            elif method.upper() in ["DELETE"]:
                test_cases += f"    response = requests.delete(url, headers=headers)\n\n"
            else:
                test_cases += f"    response = requests.request('{method}', url, headers=headers)\n\n"

            # 断言
            test_cases += f"    # 断言响应状态码为401（未授权）\n"
            test_cases += f"    assert response.status_code == 401\n"
            test_cases += f"    # 断言错误信息\n"
            test_cases += f"    assert 'error' in response.json()\n"

            test_cases += "```\n\n"

        return test_cases

    def export_test_cases(self):
        try:
            # 获取测试用例文本
            test_cases = self.test_case_preview.toPlainText()
            if not test_cases:
                QMessageBox.warning(self, "警告", "没有可以导出的测试用例")
                return

            # 获取导出格式
            export_format = self.export_format_combo.currentText()

            # 选择保存路径
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "保存测试用例",
                "",
                "JSON文件 (*.json);;YAML文件 (*.yaml);;Excel文件 (*.xlsx);;Markdown文件 (*.md);;Word文件 (*.docx)"
            )

            if not file_path:
                return

            # 根据不同格式导出
            if export_format == "JSON":
                self.export_to_json(test_cases, file_path)
            elif export_format == "YAML":
                self.export_to_yaml(test_cases, file_path)
            elif export_format == "Excel":
                self.export_to_excel(test_cases, file_path)
            elif export_format == "Markdown":
                self.export_to_markdown(test_cases, file_path)
            elif export_format == "Docx":
                self.export_to_docx(test_cases, file_path)

            QMessageBox.information(self, "成功", f"测试用例已成功导出为{export_format}格式")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出测试用例失败: {str(e)}")

    def export_to_json(self, test_cases, file_path):
        # 将Markdown格式的测试用例转换为JSON结构
        # 这里使用简单的解析逻辑，实际项目可能需要更复杂的解析

        result = {
            "test_cases": []
        }

        # 解析测试用例
        api_pattern = r'## API \d+: ([A-Z]+) (.*?)\n\n(.*?)(?=## API|\Z)'
        api_matches = re.findall(api_pattern, test_cases, re.DOTALL)

        for method, url, content in api_matches:
            api_case = {
                "method": method,
                "url": url,
                "test_cases": []
            }

            # 解析各类测试用例
            case_pattern = r'### (.*?)\n\n```python\n(.*?)```'
            case_matches = re.findall(case_pattern, content, re.DOTALL)

            for case_type, case_code in case_matches:
                api_case["test_cases"].append({
                    "type": case_type,
                    "code": case_code
                })

            result["test_cases"].append(api_case)

        # 保存为JSON文件
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

    def export_to_yaml(self, test_cases, file_path):
        # 先转换为JSON，再导出为YAML
        # 解析同export_to_json

        result = {
            "test_cases": []
        }

        # 解析测试用例
        api_pattern = r'## API \d+: ([A-Z]+) (.*?)\n\n(.*?)(?=## API|\Z)'
        api_matches = re.findall(api_pattern, test_cases, re.DOTALL)

        for method, url, content in api_matches:
            api_case = {
                "method": method,
                "url": url,
                "test_cases": []
            }

            # 解析各类测试用例
            case_pattern = r'### (.*?)\n\n```python\n(.*?)```'
            case_matches = re.findall(case_pattern, content, re.DOTALL)

            for case_type, case_code in case_matches:
                api_case["test_cases"].append({
                    "type": case_type,
                    "code": case_code
                })

            result["test_cases"].append(api_case)

        # 保存为YAML文件
        with open(file_path, 'w', encoding='utf-8') as f:
            yaml.dump(result, f, allow_unicode=True)

    def export_to_excel(self, test_cases, file_path):
        # 将测试用例导出为Excel

        # 创建Excel工作簿
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill

        wb = Workbook()
        ws = wb.active
        ws.title = "测试用例"

        # 设置表头
        headers = ["API", "方法", "URL", "测试类型", "测试代码"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")

        # 解析测试用例
        api_pattern = r'## API (\d+): ([A-Z]+) (.*?)\n\n(.*?)(?=## API|\Z)'
        api_matches = re.findall(api_pattern, test_cases, re.DOTALL)

        row = 2
        for api_num, method, url, content in api_matches:
            # 解析各类测试用例
            case_pattern = r'### (.*?)\n\n```python\n(.*?)```'
            case_matches = re.findall(case_pattern, content, re.DOTALL)

            for case_type, case_code in case_matches:
                ws.cell(row=row, column=1).value = f"API {api_num}"
                ws.cell(row=row, column=2).value = method
                ws.cell(row=row, column=3).value = url
                ws.cell(row=row, column=4).value = case_type
                ws.cell(row=row, column=5).value = case_code

                row += 1

        # 调整列宽
        for col in range(1, 6):
            ws.column_dimensions[chr(64 + col)].width = 20
        ws.column_dimensions['E'].width = 80

        # 保存Excel文件
        wb.save(file_path)

    def export_to_markdown(self, test_cases, file_path):
        # 测试用例已经是Markdown格式，直接保存
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(test_cases)

    def export_to_docx(self, test_cases, file_path, WD_ALIGN_PARAGRAPH=None):
        # 将Markdown格式的测试用例转换为Word文档
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 添加标题
        title = doc.add_heading("接口测试用例", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 转换Markdown为纯文本
        h = html2text.HTML2Text()
        h.ignore_links = False

        # 解析测试用例
        api_pattern = r'## API (\d+): ([A-Z]+) (.*?)\n\n(.*?)(?=## API|\Z)'
        api_matches = re.findall(api_pattern, test_cases, re.DOTALL)

        for api_num, method, url, content in api_matches:
            # 添加API标题
            doc.add_heading(f"API {api_num}: {method} {url}", level=1)

            # 解析各类测试用例
            case_pattern = r'### (.*?)\n\n```python\n(.*?)```'
            case_matches = re.findall(case_pattern, content, re.DOTALL)

            for case_type, case_code in case_matches:
                # 添加测试用例类型
                doc.add_heading(case_type, level=2)

                # 添加代码
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(case_code)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)

                # 添加分隔线
                doc.add_paragraph("---")

        # 保存Word文档
        doc.save(file_path)


def main():
    app = QApplication(sys.argv)
    window = MultiFormatTestTool()
    window.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()