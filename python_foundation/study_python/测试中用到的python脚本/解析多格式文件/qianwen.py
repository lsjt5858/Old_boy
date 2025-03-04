#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Author: 熊🐻来个🥬
# @Date:  2025/3/3
# @Description: [对文件功能等的简要描述（可自行添加）]
# 工具支持导入 HAR 文件、Swagger 文档（JSON/YAML 格式）、
# Excel 文件、Markdown 文档、JSON 文件、YAML 文件、TXT 文件，
# 并在 GUI 界面中预览文档内容。har包案例文件已上传云盘。

# 工具功能概述
# 多格式文档支持：
# 工具支持导入 HAR 文件、Swagger 文档（JSON/YAML 格式）、Excel 文件、Markdown 文档、JSON 文件、YAML 文件、TXT 文件，并在 GUI 界面中预览文档内容。har包案例文件已上传云盘。
#
# 文档解析与预览：
# 自动解析文档内容，提取接口信息（如请求地址、请求方法、请求头、请求参数、请求响应等），并以可视化方式展示解析结果。
#
# 智能推理生成用例：
# 结合用户输入的提示词和文档内容，调用阿里云百炼 DeepSeek API，生成智能接口测试用例脚本。
#
# 多格式导出：
# 支持将生成的接口测试用例导出为 JSON、YAML、Excel、Markdown和Docx格式，满足多种使用场景需求。


import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import json
import yaml
import pandas as pd
import requests
from docx import Document
from urllib.parse import urlparse
import markdown
import haralyzer
import threading


class APITestTool:
    def __init__(self, root):
        self.root = root
        self.root.title("智能接口测试工具 v1.0")
        self.root.geometry("1200x800")

        # 左侧面板 - 文件导入
        self.left_frame = tk.Frame(root)
        self.left_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.Y)

        self.import_btn = tk.Button(self.left_frame, text="导入文件", command=self.import_file)
        self.import_btn.pack(pady=5)

        self.format_label = tk.Label(self.left_frame, text="支持格式：HAR/Swagger/Excel/Markdown/JSON/YAML/TXT")
        self.format_label.pack(pady=5)

        # 中间面板 - 文档预览
        self.preview_frame = tk.Frame(root)
        self.preview_frame.pack(side=tk.LEFT, padx=10, pady=10, expand=True, fill=tk.BOTH)

        self.preview_label = tk.Label(self.preview_frame, text="文档预览")
        self.preview_label.pack()

        self.preview_tree = ttk.Treeview(self.preview_frame, columns=("method", "url"), show="tree")
        self.preview_tree.column("#0", width=200)
        self.preview_tree.column("method", width=100)
        self.preview_tree.column("url", width=400)
        self.preview_tree.pack(fill=tk.BOTH, expand=True)
        self.preview_tree.bind("<<TreeviewSelect>>", self.show_details)

        self.detail_text = tk.Text(self.preview_frame, height=10)
        self.detail_text.pack(fill=tk.X)

        # 右侧面板 - 用例生成
        self.right_frame = tk.Frame(root)
        self.right_frame.pack(side=tk.RIGHT, padx=10, pady=10, fill=tk.Y)

        self.prompt_label = tk.Label(self.right_frame, text="输入提示词：")
        self.prompt_label.pack()
        self.prompt_entry = tk.Entry(self.right_frame, width=30)
        self.prompt_entry.pack()

        self.api_key_label = tk.Label(self.right_frame, text="API密钥：")
        self.api_key_label.pack()
        self.api_key_entry = tk.Entry(self.right_frame, width=30, show="*")
        self.api_key_entry.pack()

        self.generate_btn = tk.Button(self.right_frame, text="生成用例", command=self.generate_test_case)
        self.generate_btn.pack(pady=10)

        self.export_btn = tk.Button(self.right_frame, text="导出用例", command=self.export_test_case)
        self.export_btn.pack(pady=5)

        # 状态栏
        self.status = tk.Label(root, text="就绪", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status.pack(side=tk.BOTTOM, fill=tk.X)

        # 数据存储
        self.endpoints = []
        self.current_data = None
        self.test_cases = []

    def import_file(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("All Supported", "*.har *.json *.yaml *.yml *.xlsx *.md *.txt"),
            ("HAR", "*.har"), ("Swagger", "*.json *.yaml *.yml"),
            ("Excel", "*.xlsx"), ("Markdown", "*.md"), ("Text", "*.txt")])

        if not file_path:
            return

        self.status.config(text=f"导入文件: {file_path}")
        threading.Thread(target=self.parse_file, args=(file_path,)).start()

    def parse_file(self, file_path):
        try:
            if file_path.endswith('.har'):
                with open(file_path, 'r') as f:
                    har_data = json.load(f)
                self.parse_har(har_data)

            elif file_path.endswith(('.json', '.yaml', '.yml')):
                self.parse_swagger(file_path)

            elif file_path.endswith('.xlsx'):
                self.parse_excel(file_path)

            elif file_path.endswith('.md'):
                self.parse_markdown(file_path)

            elif file_path.endswith('.txt'):
                self.parse_text(file_path)

            self.root.after(0, self.update_preview)
            self.status.config(text="解析完成")

        except Exception as e:
            self.status.config(text=f"解析失败: {str(e)}")
            messagebox.showerror("错误", f"解析失败：{str(e)}")

    def parse_har(self, har_data):
        har_parser = haralyzer.HarParser(har_data)
        self.endpoints = []
        for entry in har_parser.har_data['entries']:
            request = entry['request']
            url = request['url']
            method = request['method']
            headers = {h['name']: h['value'] for h in request['headers']}
            query = {q['name']: q['value'] for q in request.get('queryString', [])}
            post_data = request.get('postData', {}).get('text', '')

            self.endpoints.append({
                'method': method,
                'url': url,
                'headers': headers,
                'query_params': query,
                'body': post_data
            })

    def parse_swagger(self, file_path):
        with open(file_path, 'r') as f:
            if file_path.endswith('.json'):
                data = json.load(f)
            else:
                data = yaml.safe_load(f)

        self.endpoints = []
        for path, methods in data['paths'].items():
            for method, spec in methods.items():
                endpoint = {
                    'method': method.upper(),
                    'url': path,
                    'summary': spec.get('summary', ''),
                    'parameters': [],
                    'responses': spec.get('responses', {})
                }

                for param in spec.get('parameters', []):
                    endpoint['parameters'].append({
                        'name': param['name'],
                        'in': param['in'],
                        'type': param['schema']['type']
                    })

                self.endpoints.append(endpoint)

    def parse_excel(self, file_path):
        df = pd.read_excel(file_path)
        self.endpoints = df.to_dict('records')

    def parse_markdown(self, file_path):
        with open(file_path, 'r') as f:
            md = markdown.Markdown(extensions=['tables'])
            html = md.convert(f.read())
        # 这里需要根据实际markdown格式实现解析逻辑
        # 示例：提取表格数据
        tables = md.tables
        if tables:
            self.endpoints = tables[0].to_dict('records')

    def parse_text(self, file_path):
        with open(file_path, 'r') as f:
            lines = f.readlines()
        # 示例：按行解析简单文本格式
        self.endpoints = [{'endpoint': line.strip()} for line in lines]

    def update_preview(self):
        self.preview_tree.delete(*self.preview_tree.get_children())
        for endpoint in self.endpoints:
            method = endpoint.get('method', 'N/A')
            url = endpoint.get('url', 'N/A')
            self.preview_tree.insert("", "end", text=urlparse(url).path,
                                     values=(method, url))

    def show_details(self, event):
        selected_item = self.preview_tree.selection()[0]
        item = self.preview_tree.item(selected_item)
        url = item['values'][1]
        endpoint = next((e for e in self.endpoints if e['url'] == url), {})

        self.detail_text.delete(1.0, tk.END)
        self.detail_text.insert(tk.END, json.dumps(endpoint, indent=2))

    def generate_test_case(self):
        prompt = self.prompt_entry.get()
        api_key = self.api_key_entry.get()
        if not api_key:
            messagebox.showwarning("警告", "请输入API密钥")
            return

        # 调用DeepSeek API
        try:
            response = requests.post(
                "https://api.deepseek.com/v1/casegen",
                headers={"Authorization": f"Bearer {api_key}"},
                json={
                    "prompt": prompt,
                    "endpoints": self.endpoints
                }
            )
            response.raise_for_status()
            self.test_cases = response.json()['test_cases']
            self.status.config(text=f"生成{len(self.test_cases)}个测试用例")
        except Exception as e:
            self.status.config(text=f"生成失败: {str(e)}")
            messagebox.showerror("错误", f"生成失败：{str(e)}")

    def export_test_case(self):
        if not self.test_cases:
            messagebox.showwarning("警告", "请先生成测试用例")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON", "*.json"), ("YAML", "*.yaml"),
                       ("Excel", "*.xlsx"), ("Markdown", "*.md"),
                       ("Word", "*.docx")]
        )

        if not file_path:
            return

        try:
            if file_path.endswith('.json'):
                with open(file_path, 'w') as f:
                    json.dump(self.test_cases, f, indent=2)

            elif file_path.endswith(('.yaml', '.yml')):
                with open(file_path, 'w') as f:
                    yaml.dump(self.test_cases, f)

            elif file_path.endswith('.xlsx'):
                df = pd.DataFrame(self.test_cases)
                df.to_excel(file_path, index=False)

            elif file_path.endswith('.md'):
                with open(file_path, 'w') as f:
                    f.write("# 测试用例\n\n")
                    for case in self.test_cases:
                        f.write(f"## {case['name']}\n")
                        f.write(f"**描述**: {case['description']}\n")
                        f.write("```json\n")
                        f.write(json.dumps(case, indent=2))
                        f.write("\n```\n")

            elif file_path.endswith('.docx'):
                doc = Document()
                doc.add_heading('测试用例', 0)
                for case in self.test_cases:
                    doc.add_heading(case['name'], level=1)
                    doc.add_paragraph(f"描述: {case['description']}")
                    doc.add_paragraph("详细信息:").style = 'List Bullet'
                    doc.add_paragraph(json.dumps(case, indent=2))
                doc.save(file_path)

            self.status.config(text=f"导出成功: {file_path}")
        except Exception as e:
            self.status.config(text=f"导出失败: {str(e)}")
            messagebox.showerror("错误", f"导出失败：{str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    tool = APITestTool(root)
    root.mainloop()


